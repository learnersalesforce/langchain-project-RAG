import fitz  # PyMuPDF for PDF processing
from docx import Document  # DOCX processing
import spacy  # NLP for entity recognition
import re  # Regex for pattern extraction
import os

# Load NLP model
nlp = spacy.load("en_core_web_sm")

# Extract Text from PDF
def extract_text_from_pdf(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = "\n".join([page.get_text("text") for page in doc])
        return text.strip() if text.strip() else "No text found in PDF"
    except Exception as e:
        print(f" Error extracting text from PDF {pdf_path}: {e}")
        return ""

# Extract Text from DOCX
def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip() if text.strip() else "No text found in DOCX"
    except Exception as e:
        print(f" Error extracting text from DOCX {docx_path}: {e}")
        return ""

#  Extract Certifications
def extract_certifications(text):
    certifications = re.findall(r"((?:Salesforce|AWS|Azure|Google Cloud|Microsoft)[\w\s]+Certified[\w\s]*)", text, re.IGNORECASE)
    return [cert.strip() for cert in certifications]

#  Extract Technical Skills
def extract_technical_skills(text):
    skills = re.findall(
        r"(JavaScript|Python|Java|C\+\+|C#|SQL|MySQL|PostgreSQL|MongoDB|Apex|SOQL|"
        r"SOSL|Visualforce|LWC|Aura|Salesforce|MuleSoft|REST API|SOAP API|GitHub|JIRA|"
        r"Data Loader|Lightning Web Components|Bitbucket|AWS|Azure|Google Cloud|Docker|Kubernetes)",
        text, re.IGNORECASE
    )
    return list(set([skill.strip() for skill in skills]))

#  Extract Projects and URLs
def extract_projects(text):
    projects = []
    project_matches = re.findall(r"(?i)Projects?\s*[:\-]?\s*(.*?)\n", text)
    url_matches = re.findall(r"https?://[^\s]+", text)

    for i, project in enumerate(project_matches):
        project_name = project.strip()
        if len(project_name.split()) > 1:
            project_url = url_matches[i] if i < len(url_matches) else "Not Available"
            projects.append(f"{project_name} ({project_url})")

    return projects

#  Extract Education (Degree, University, CGPA, Percentage)
def extract_education(text):
    education = []
    education_matches = re.findall(
        r"((?:Bachelor|Master|B\.Tech|M\.Tech|MBA|PhD|Graduate)[^,]*)\s*(?:from|at)?\s*([\w\s]+University|Institute)[,\s]*(?:CGPA[:\-]?\s*(\d+\.\d+)|Percentage[:\-]?\s*(\d+%))?",
        text, re.IGNORECASE
    )

    for match in education_matches:
        degree, university, cgpa, percentage = match
        details = f"{degree} from {university}"
        if cgpa:
            details += f" (CGPA: {cgpa})"
        if percentage:
            details += f" (Percentage: {percentage})"
        education.append(details)

    return education

#  Extract Resume Information
def extract_information(text):
    doc = nlp(text)

    name = next((ent.text for ent in doc.ents if ent.label_ == "PERSON"), "Unknown")
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    phone_match = re.search(r"\+?\d[\d\s\-()]{8,}\d", text)
    linkedin_match = re.search(r"(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9-]+", text)

    extracted_data = {
        "Name": name,
        "Email": email_match.group() if email_match else "Not Found",
        "Phone": phone_match.group() if phone_match else "Not Found",
        "LinkedIn": linkedin_match.group() if linkedin_match else "Not Found",
        "Education": ", ".join(extract_education(text)),
        "Certifications": ", ".join(extract_certifications(text)),
        "Technical Skills": ", ".join(extract_technical_skills(text)),
        "Projects": ", ".join(extract_projects(text)),
    }

    return extracted_data

#  Process a Single Resume
def process_resume(file_path):
    ext = file_path.split(".")[-1].lower()
    text = extract_text_from_pdf(file_path) if ext == "pdf" else extract_text_from_docx(file_path)

    if not text.strip():
        print(f" Error: No text extracted from {file_path}")
        return None
    
    extracted_info = extract_information(text)
    extracted_info["Resume File"] = os.path.basename(file_path)
    
    return extracted_info

#  Save Extracted Data to a DOCX Table with the Candidate's Name
from docx.shared import Pt  # Import for font size
from docx.oxml.ns import qn  # Import for XML styling
from docx.oxml import OxmlElement  # Import for modifying XML structure

from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def save_to_docx(data, folder_path):
    # Get the first valid name from extracted data
    first_valid_name = next((d["Name"] for d in data if d["Name"] != "Unknown"), "Extracted_Resumes")

    # Generate DOCX file with the extracted name
    doc_filename = f"{first_valid_name.replace(' ', '_')}-Resume.docx"
    doc_path = os.path.join(folder_path, doc_filename)

    doc = Document()
    doc.add_heading("Extracted Resume Data", level=1)

    for resume in data:
        doc.add_heading(f"Resume: {resume['Resume File']}", level=2)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Set table header and apply bold formatting
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Field"
        hdr_cells[1].text = "Value"
        
        for cell in hdr_cells:
            run = cell.paragraphs[0].runs[0]
            run.bold = True  # Make header bold
            run.font.size = Pt(11)

        # Add extracted data rows with bold field names
        for key, value in resume.items():
            row_cells = table.add_row().cells

            # Field Name (Bold)
            field_run = row_cells[0].paragraphs[0].add_run(key)
            field_run.bold = True  # Bold the field name
            field_run.font.size = Pt(11)  # Set font size

            # Field Value
            row_cells[1].text = str(value)

        doc.add_paragraph("\n")  # Add spacing

    doc.save(doc_path)
    print(f"\n Data saved to {doc_path}")



#  Run the Script
if __name__ == "__main__":
    folder_path = "D:/Updated-Langchain/read_resume"
    extracted_data = []

    for file in os.listdir(folder_path):
        if file.lower().endswith((".pdf", ".docx")):
            file_path = os.path.join(folder_path, file)
            print(f"\n🚀 Processing: {file_path}")
            data = process_resume(file_path)
            if data:
                extracted_data.append(data)

    if extracted_data:
        save_to_docx(extracted_data, folder_path)
